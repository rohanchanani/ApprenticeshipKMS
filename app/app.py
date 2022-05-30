import docx
import docx2txt
import sqlalchemy
import pandas as pd
from flask import Flask, render_template, request, redirect
import os
import sys
import json
import pymongo
from pymongo import MongoClient
import re

def get_database():

    # Provide the mongodb atlas url to connect python to mongodb using pymongo
    CONNECTION_STRING = "mongodb+srv://apprenticeship:" + os.environ["mongo_password"] + "@kmsdataextraction.jodxa.mongodb.net/?retryWrites=true&w=majority"

    # Create a connection using MongoClient. You can import MongoClient or use pymongo.MongoClient
    client = MongoClient(CONNECTION_STRING)

    # Create the database for our example (we will use the same database throughout the tutorial
    return client['data-extraction']

extraction_db = get_database()

companies = extraction_db.companies
questions = extraction_db.questions

app = Flask(__name__)

def upload_question_to_database(question, answer, section, company, version, number, question_results, company_results):   
    if not question_results.get(question+answer):
        questions.insert_one({"number":number, "question":question, "answer":answer, "version": version, "company":company, "section":section})

def read_questions(company, version, current_section, i, all_lines, found_questions, sections):
    question_results = {test["question"]+test["answer"]:1 for test in list(questions.find({}, {"_id":0}))}
    company_results = {test["company_name"]:1 for test in list(companies.find({}, {"_id":0}))}
    question_number = 1
    while i < len(all_lines):
        current_question = all_lines[i]
        try:
            next_ten_questions = found_questions[found_questions.index(current_question) + 1: found_questions.index(current_question) + 11]
        except:
            i += 1
            continue
        current_answer = all_lines[i+1]
        i += 2
        while i < len(all_lines):
            if all_lines[i] in next_ten_questions:
                break
            if all_lines[i] == sections[min(sections.index(current_section) + 1, len(sections) - 1)]:
                current_section = all_lines[i]
                i += 1
                break
            current_answer += "\n" + all_lines[i]
            i += 1
        upload_question_to_database(current_question, current_answer, current_section, company, version, question_number, question_results, company_results)
        question_number += 1

@app.before_request
def beforeRequest():
    if not request.url.startswith('https') and 'DYNO' in os.environ:
        return redirect(request.url.replace('http', 'https', 1))

@app.route("/query", methods=["POST"])
def sql_query():
    query = {}
    for column in request.form.keys():
        if request.form[column].strip() != "":
            if column == "number":
                query["number"] = int(request.form["number"])
                continue
            rgx = re.compile(".*" + request.form[column]+".*", re.IGNORECASE)
            query[column] = rgx
    print(query)
    try:
        result = list(questions.find(query, {"_id":0}))
    except:
        return "Please enter a valid query"
    columns = questions.find_one({}, {"_id":0}).keys()
    return render_template("query.html", columns=columns, rows=result) 

@app.route("/company", methods=["POST"])
def get_company():
    company_name = request.form.get("company")
    current_questions = list(questions.find({"company":company_name}, {"_id":0}))
    return render_template("company.html", questions=current_questions)

@app.route("/", methods=["GET", "POST"])
def homepage():
    if request.method == "GET":
        current_questions = questions.find({}, {"_id":0})
        try:
            question_columns = questions.find_one({}, {"_id":0}).keys()
        except:
            question_columns = ["there", "are", "no", "columns", "yet"]
        current_companies = list(companies.find({}, {"_id":0}))
        return render_template("index.html", questions=list(current_questions), companies=current_companies, columns=question_columns)
    else:
        for filename, file in request.files.items():
            file.save(filename)
            text = docx2txt.process(filename)
            all_lines = [line.strip().replace("'", "") for line in text.split("\n") if line != ""]
            with open(filename, 'w') as f:
                pass
            doc = docx.Document(file)
            found_questions = [paragraph.text.strip().replace("'", "") for paragraph in doc.paragraphs if paragraph.text != "" and paragraph.style.name in ["Heading 2", "Heading 3"]]
            sections = [paragraph.text.strip().replace("'", "") for paragraph in doc.paragraphs if paragraph.text != "" and paragraph.style.name == "Heading 1"]
            i = 0
            version_not_found = True
            current_section = ""
            questions_started = False
            while not questions_started:
                if all_lines[i] == "Technology Assessment of":
                    company_name = all_lines[i+1]
                    if not len(list(companies.find({"company_name":company_name}))):
                        companies.insert_one({"company_name":company_name})
                    i += 2
                elif "Document Version" in all_lines[i] and version_not_found:
                    document_version = all_lines[i].split(" ")[-1]
                    version_not_found = False
                    i += 1
                elif all_lines[i] in sections:
                    current_section = all_lines[i]
                    i += 1
                    questions_started = True
                else:
                    i += 1

            read_questions(company_name, document_version, current_section, i, all_lines, found_questions, sections)
        current_questions = questions.find({}, {"_id":0})
        try:
            question_columns = questions.find_one({}, {"_id":0}).keys()
        except:
            question_columns = ["there", "are", "no", "columns", "yet"]
        current_companies = list(companies.find({}, {"_id":0}))
        return render_template("index.html", questions=current_questions, companies=current_companies, columns=question_columns)